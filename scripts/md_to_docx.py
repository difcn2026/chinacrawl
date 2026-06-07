"""
Xiao Hei XHLS v3.0 - Markdown to WPS/DOCX Batch Converter
Usage: python md_to_docx.py <input_dir> <output_dir>
Converts .md files in input_dir to .docx, openable with WPS/Word
"""
import os, sys, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.ns import qn
from lxml import etree
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def md_to_docx(md_path, docx_path):
    """Convert a single markdown file to docx."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Microsoft YaHei"
        heading_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    i = 0
    in_code_block = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(1)
            i += 1
            continue

        if not line.strip():
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        if line.strip() in ("---", "***"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = etree.SubElement(pPr, qn("w:pBdr"))
            bottom = etree.SubElement(pBdr, qn("w:bottom"))
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if line.strip().startswith("|---"):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table:
            if table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False

        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue
        elif line.startswith("## ") and not line.startswith("### "):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue

        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1

    if in_table and table_rows:
        _add_table(doc, table_rows)

    doc.save(docx_path)
    return True


def _add_table(doc, rows):
    """Add a table to the document from parsed markdown rows."""
    if not rows or len(rows) < 1:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]), style="Light Grid Accent 1")
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < len(table.rows[r_idx].cells):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = cell_text
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
    doc.add_paragraph()


def _add_formatted_text(paragraph, text):
    """Add text to paragraph handling bold and code inline."""
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xcc, 0x33, 0x33)
        else:
            paragraph.add_run(part)


def batch_convert(input_dir, output_dir):
    """Convert all .md files in input_dir to .docx in output_dir."""
    os.makedirs(output_dir, exist_ok=True)

    md_files = [f for f in os.listdir(input_dir) if f.endswith(".md")]
    if not md_files:
        print("No .md files found.")
        return

    for md_file in md_files:
        md_path = os.path.join(input_dir, md_file)
        docx_name = md_file.replace(".md", ".docx")
        docx_path = os.path.join(output_dir, docx_name)
        try:
            md_to_docx(md_path, docx_path)
            print(f"  [OK] {md_file} -> {docx_name}")
        except Exception as e:
            print(f"  [FAIL] {md_file}: {e}")

    print(f"\nConverted {len(md_files)} files -> {output_dir}")


if __name__ == "__main__":
    if len(sys.argv) >= 3:
        batch_convert(sys.argv[1], sys.argv[2])
    else:
        base = r"C:\Users\Administrator\Documents\New project"
        input_dir = os.path.join(base, "knowledge", "pipeline-decisions")
        output_dir = os.path.join(base, "docs", "wps")
        print(f"Input:  {input_dir}")
        print(f"Output: {output_dir}")
        print()
        batch_convert(input_dir, output_dir)
